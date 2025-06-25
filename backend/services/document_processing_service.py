import os
import json
import time
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional
import fitz
from docx import Document as DocxDocument
import olefile
import zipfile


import google.generativeai as genai
import sys

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import GEMINI_API_KEY, DATA_DIR

class DocumentProcessingService:
    def __init__(self):
        """Khởi tạo dịch vụ xử lý tài liệu"""
        genai.configure(api_key=GEMINI_API_KEY)
        print("Dịch vụ xử lý tài liệu đã được khởi tạo - hỗ trợ PDF và Word")
        
        # Template prompt cho Gemini
        self.chunking_prompt = """
Bạn là chuyên gia phân tích và chia nhỏ văn bản pháp luật Việt Nam. 
Nhiệm vụ của bạn là chia văn bản thành các chunk hợp lý và trích xuất thông tin về các văn bản liên quan.

1. **Nguyên tắc chia chunk:**
   - Mỗi chunk khoảng 500-1000 từ
   - Ưu tiên chia theo điều, khoản, mục
   - Nhóm các khoản liên quan trong cùng một điều
   - Tách riêng phần mở đầu (preamble) bao gồm tiêu đề, căn cứ
   - Tách riêng phụ lục, biểu mẫu nếu có

2. **Quy tắc đặt tên chunk_id:**
   - Format: {doc_id}_{loại}
   - Loại: preamble, art[số], appendix, form
   - Ví dụ: 47_2009_TTLT_BTC_BLĐTBXH_preamble, 47_2009_TTLT_BTC_BLĐTBXH_art1_2

3. **Loại chunk_type:**
   - preamble: Phần mở đầu, căn cứ
   - article: Các điều khoản chính
   - appendix: Phụ lục
   - form: Biểu mẫu

4. **Trích xuất thông tin văn bản liên quan (related_documents):**
   Tìm và phân tích các thông tin sau trong văn bản:
   
   a) **Căn cứ pháp lý** (relationship: "references"):
      - Tìm trong phần "Căn cứ" ở đầu văn bản
      - Các văn bản được tham chiếu: Luật, Nghị định, Thông tư...
      - Format: số_năm_loại_cơquan (ví dụ: 26_2013_NĐ_CP)
   
   b) **Văn bản thay thế** (relationship: "replaces"):
      - Tìm cụm từ: "thay thế", "thay thế cho"
      - Thường ở điều cuối hoặc phần hiệu lực
   
   c) **Văn bản được thay thế bởi** (relationship: "replaced_by"):
      - Nếu có thông tin văn bản này thay thế văn bản khác
   
   d) **Văn bản sửa đổi** (relationship: "amends"):
      - Tìm cụm từ: "sửa đổi", "bổ sung"
      - Văn bản nào được sửa đổi/bổ sung
   
   e) **Văn bản hướng dẫn** (relationship: "implements"):
      - Tìm cụm từ: "hướng dẫn thi hành", "quy định chi tiết"
      - Văn bản cấp trên được hướng dẫn

5. **Trích xuất thông tin ngày tháng:**
   - issue_date: Ngày ban hành (tìm trong phần đầu văn bản)
   - effective_date: Ngày có hiệu lực (thường ở điều cuối)
   - expiry_date: Ngày hết hiệu lực (nếu có)

6. **QUAN TRỌNG - Tự động detect metadata:**
   - doc_id: Trích xuất từ tiêu đề hoặc phần đầu văn bản (format: số_năm_loại_cơquan)
   - doc_type: Xác định loại văn bản (Luật, Nghị định, Thông tư, Quyết định, Pháp lệnh)
   - doc_title: Trích xuất tiêu đề đầy đủ của văn bản
   - effective_date: Tìm ngày có hiệu lực trong văn bản

7. **Yêu cầu output:**
   - Trả về JSON với cấu trúc metadata hoàn chỉnh
   - Giữ nguyên 100% nội dung gốc trong các chunk
   - Đưa ra content_summary mô tả ngắn gọn nội dung chunk
   - Trích xuất chính xác thông tin related_documents
   - TỰ ĐỘNG ĐIỀN đầy đủ thông tin metadata

**LƯU Ý QUAN TRỌNG:**
- Đọc kỹ toàn bộ văn bản để tìm thông tin về văn bản liên quan
- Chú ý các cụm từ tiếng Việt: "Căn cứ", "Thay thế", "Sửa đổi", "Bổ sung", "Hướng dẫn thi hành"
- Trích xuất chính xác số hiệu văn bản theo format: số_năm_loại_cơquan
- Nếu không tìm thấy thông tin nào, để array rỗng []
- PHẢI tự động detect và điền đầy đủ metadata từ nội dung văn bản

Văn bản cần phân tích:
{content}

Thông tin cơ bản ban đầu (có thể điều chỉnh dựa trên nội dung):
- doc_id: {doc_id}
- doc_type: {doc_type}
- doc_title: {doc_title}
- effective_date: {effective_date}
- document_scope: {document_scope}

Trả về JSON theo format sau (chú ý phân tích kỹ để điền đầy đủ metadata được AUTO-DETECT):
{{
  "doc_id": "auto_detected_doc_id_hoặc_{doc_id}",
  "doc_type": "auto_detected_doc_type_hoặc_{doc_type}",
  "doc_title": "auto_detected_title_hoặc_{doc_title}",
  "issue_date": "ngày_ban_hành_từ_văn_bản (DD-MM-YYYY)",
  "effective_date": "ngày_hiệu_lực_từ_văn_bản_hoặc_{effective_date}",
  "expiry_date": "ngày_hết_hiệu_lực_nếu_có_hoặc_null",
  "status": "active",
  "document_scope": "{document_scope}",
  "replaces": ["danh_sách_văn_bản_bị_thay_thế"],
  "replaced_by": "văn_bản_thay_thế_hoặc_null",
  "amends": "văn_bản_được_sửa_đổi_hoặc_null",
  "amended_by": "văn_bản_sửa_đổi_hoặc_null",
  "retroactive": false,
  "retroactive_date": null,
  "chunks": [
    {{
      "chunk_id": "chunk_id_theo_quy_tac",
      "chunk_type": "loai_chunk",
      "file_path": "/data/{doc_id}/chunk_X.md",
      "content_summary": "mo_ta_noi_dung",
      "content": "noi_dung_chunk_day_du"
    }}
  ],
  "related_documents": [
    {{
      "doc_id": "số_năm_loại_cơquan",
      "relationship": "references|replaces|replaced_by|amends|amended_by|implements",
      "description": "mô_tả_chi_tiết_về_mối_quan_hệ"
    }}
  ]
}}

**VÍ DỤ VỀ AUTO-DETECTION:**
Nếu văn bản có tiêu đề "THÔNG TƯ 47/2009/TTLT-BTC-BLĐTBXH", thì:
- doc_id: "47_2009_TTLT_BTC_BLĐTBXH"
- doc_type: "Thông tư"
- doc_title: trích xuất tiêu đề đầy đủ từ văn bản

LƯU Ý: NẾU NGƯỜI DÙNG CUNG CẤP VĂN BẢN KHÔNG THUỘC VỀ LĨNH VỰC PHÁP LUẬT VIỆT NAM, HÃY TRẢ VỀ MỘT JSON RỖNG VỚI CÁC TRƯỜNG BẮT BUỘC."""

    def extract_pdf_content(self, file_path: str) -> str:
        """Trích xuất nội dung từ file PDF"""
        if not fitz:
            raise Exception("PyMuPDF không được cài đặt. Cần cài: pip install PyMuPDF")
            
        try:
            print(f"Đang trích xuất nội dung từ PDF: {file_path}")
            
            doc = fitz.open(file_path)
            full_text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                full_text += text + "\n"
                print(f"Đã trích xuất trang {page_num + 1}/{len(doc)}")
            
            doc.close()
            
            # Làm sạch text nhưng giữ cấu trúc
            full_text = full_text.replace('\n\n\n', '\n\n')
            full_text = full_text.strip()
            
            print(f"Hoàn thành trích xuất PDF. Tổng số ký tự: {len(full_text)}")
            return full_text
            
        except Exception as e:
            print(f"Lỗi khi trích xuất PDF: {str(e)}")
            raise Exception(f"Không thể đọc file PDF: {str(e)}")

    def extract_docx_content(self, file_path: str) -> str:
        try:
            print(f"Đang trích xuất nội dung từ Word DOCX: {file_path}")
            
            doc = DocxDocument(file_path)
            full_text = ""
            
            # Trích xuất text từ các paragraph
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"
            
            # Trích xuất text từ tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text += " | ".join(row_text) + "\n"
            
            # Làm sạch text
            full_text = full_text.replace('\n\n\n', '\n\n')
            full_text = full_text.strip()
            
            print(f"Hoàn thành trích xuất DOCX. Tổng số ký tự: {len(full_text)}")
            return full_text
            
        except Exception as e:
            print(f"Lỗi khi trích xuất DOCX: {str(e)}")
            raise Exception(f"Không thể đọc file Word DOCX: {str(e)}")

    def extract_doc_content(self, file_path: str) -> str:
        """Trích xuất nội dung từ file Word .doc (legacy format)"""
        try:
            print(f"Đang trích xuất nội dung từ Word DOC: {file_path}")
            
            # Thử convert DOC sang text bằng cách đọc binary và tìm text
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Simple text extraction - có thể không hoàn hảo với DOC format
            text_content = ""
            try:
                # Thử decode như text
                decoded = content.decode('latin-1', errors='ignore')
                # Lọc ra các ký tự có thể đọc được
                text_content = ''.join(char for char in decoded if char.isprintable() or char.isspace())
                
                # Làm sạch text
                lines = text_content.split('\n')
                cleaned_lines = []
                for line in lines:
                    clean_line = line.strip()
                    if len(clean_line) > 3:  # Bỏ qua dòng quá ngắn
                        cleaned_lines.append(clean_line)
                
                text_content = '\n'.join(cleaned_lines)
                
            except Exception as e:
                print(f"Lỗi khi decode DOC content: {str(e)}")
                raise Exception("Không thể đọc file Word DOC. Vui lòng convert sang DOCX hoặc PDF.")
            
            print(f"Hoàn thành trích xuất DOC. Tổng số ký tự: {len(text_content)}")
            return text_content
            
        except Exception as e:
            print(f"Lỗi khi trích xuất DOC: {str(e)}")
            raise Exception(f"Không thể đọc file Word DOC: {str(e)}")

    def extract_document_content(self, file_path: str) -> str:
        """Trích xuất nội dung từ file tài liệu (PDF, DOCX, DOC)"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        print(f"Detecting file type: {file_extension}")
        
        if file_extension == '.pdf':
            return self.extract_pdf_content(file_path)
        elif file_extension == '.docx':
            return self.extract_docx_content(file_path)
        elif file_extension == '.doc':
            return self.extract_doc_content(file_path)
        else:
            raise Exception(f"Định dạng file không được hỗ trợ: {file_extension}. Chỉ hỗ trợ PDF, DOCX, DOC")

    def chunk_content_with_gemini(self, content: str, doc_metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Sử dụng Gemini để chia chunk văn bản và auto-detect metadata"""
        try:
            print("Đang gọi Gemini để phân tích văn bản, chia chunk và auto-detect metadata...")
            
            # Tạo prompt với hướng dẫn chi tiết về auto-detection
            prompt = self.chunking_prompt.format(
                content=content,
                doc_id=doc_metadata.get('doc_id', ''),
                doc_type=doc_metadata.get('doc_type', ''),
                doc_title=doc_metadata.get('doc_title', ''),
                effective_date=doc_metadata.get('effective_date', ''),
                document_scope=doc_metadata.get('document_scope', 'Quốc gia')
            )
            
            print("Đã tạo prompt để phân tích văn bản với auto-detection")
            
            # Gọi Gemini API với temperature thấp để có kết quả ổn định
            model = genai.GenerativeModel('gemini-2.0-flash')
            response = model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.1,  # Giảm tính ngẫu nhiên
                    top_p=0.8,
                    top_k=40,
                )
            )
            
            print("Đã nhận phản hồi từ Gemini với auto-detection")
            
            # Parse JSON response
            response_text = response.text.strip()
            print("Preview phản hồi từ Gemini:")
            print(response_text[:300] + "..." if len(response_text) > 300 else response_text)
            
            # Tìm JSON trong response
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            
            if json_start == -1 or json_end <= json_start:
                raise ValueError("Không tìm thấy JSON hợp lệ trong phản hồi của Gemini")
            
            json_str = response_text[json_start:json_end]
            result = json.loads(json_str)
            
            chunks_count = len(result.get('chunks', []))
            related_docs_count = len(result.get('related_documents', []))
            
            print(f"Gemini đã phân tích thành công:")
            print(f"  - Auto-detected doc_id: {result.get('doc_id', 'N/A')}")
            print(f"  - Auto-detected doc_type: {result.get('doc_type', 'N/A')}")
            print(f"  - Auto-detected doc_title: {result.get('doc_title', 'N/A')[:50]}...")
            print(f"  - Auto-detected effective_date: {result.get('effective_date', 'N/A')}")
            print(f"  - Chia thành {chunks_count} chunks")
            print(f"  - Tìm thấy {related_docs_count} văn bản liên quan")
            
            # Log thông tin related_documents
            if related_docs_count > 0:
                print("Chi tiết văn bản liên quan:")
                for doc in result.get('related_documents', []):
                    print(f"  - {doc.get('doc_id', 'N/A')} ({doc.get('relationship', 'N/A')}): {doc.get('description', 'N/A')[:50]}...")
            
            # Validate và làm sạch dữ liệu
            result = self._validate_and_clean_result(result, doc_metadata)
            
            return result
            
        except json.JSONDecodeError as je:
            print(f"Lỗi parse JSON từ Gemini: {str(je)}")
            print("Raw response từ Gemini:")
            print(response_text)
            raise Exception(f"Gemini trả về JSON không hợp lệ: {str(je)}")
        except Exception as e:
            print(f"Lỗi khi gọi Gemini: {str(e)}")
            raise Exception(f"Lỗi xử lý với Gemini: {str(e)}")

    def _validate_and_clean_result(self, result: Dict[str, Any], original_metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Validate và làm sạch kết quả từ Gemini với auto-detection"""
        try:
            print("Đang validate và làm sạch kết quả từ Gemini với auto-detection...")
            
            # Ưu tiên dữ liệu auto-detected từ Gemini, fallback về original nếu cần
            validated_result = {
                "doc_id": result.get("doc_id") or original_metadata.get('doc_id', ''),
                "doc_type": result.get("doc_type") or original_metadata.get('doc_type', ''),
                "doc_title": result.get("doc_title") or original_metadata.get('doc_title', ''),
                "issue_date": result.get("issue_date", datetime.now().strftime("%d-%m-%Y")),
                "effective_date": result.get("effective_date") or original_metadata.get('effective_date', ''),
                "expiry_date": result.get("expiry_date"),
                "status": result.get("status", "active"),
                "document_scope": result.get("document_scope") or original_metadata.get('document_scope', 'Quốc gia'),
                "replaces": result.get("replaces", []),
                "replaced_by": result.get("replaced_by"),
                "amends": result.get("amends"),
                "amended_by": result.get("amended_by"),
                "retroactive": result.get("retroactive", False),
                "retroactive_date": result.get("retroactive_date"),
                "chunks": result.get("chunks", []),
                "related_documents": result.get("related_documents", [])
            }
            
            print(f"Auto-detection results:")
            print(f"  - doc_id: {validated_result['doc_id']}")
            print(f"  - doc_type: {validated_result['doc_type']}")
            print(f"  - doc_title: {validated_result['doc_title'][:50]}...")
            print(f"  - effective_date: {validated_result['effective_date']}")
            
            # Validate related_documents
            cleaned_related_docs = []
            valid_relationships = ["references", "replaces", "replaced_by", "amends", "amended_by", "implements"]
            
            for doc in validated_result["related_documents"]:
                if isinstance(doc, dict) and "doc_id" in doc and "relationship" in doc:
                    # Làm sạch doc_id (loại bỏ ký tự đặc biệt)
                    doc_id = str(doc["doc_id"]).strip()
                    relationship = str(doc["relationship"]).strip().lower()
                    
                    # Validate relationship
                    if relationship in valid_relationships:
                        cleaned_doc = {
                            "doc_id": doc_id,
                            "relationship": relationship,
                            "description": str(doc.get("description", "")).strip()
                        }
                        cleaned_related_docs.append(cleaned_doc)
                        print(f"Validated related document: {doc_id} ({relationship})")
                    else:
                        print(f"Bỏ qua relationship không hợp lệ: {relationship}")
                else:
                    print(f"Bỏ qua related document không hợp lệ: {doc}")
            
            validated_result["related_documents"] = cleaned_related_docs
            
            # Validate chunks
            if not validated_result["chunks"]:
                raise ValueError("Không có chunk nào được tạo")
            
            print(f"Validation hoàn thành:")
            print(f"  - {len(validated_result['chunks'])} chunks hợp lệ")
            print(f"  - {len(cleaned_related_docs)} văn bản liên quan hợp lệ")
            
            return validated_result
            
        except Exception as e:
            print(f"Lỗi trong quá trình validate: {str(e)}")
            raise Exception(f"Lỗi validate kết quả: {str(e)}")

    def save_chunks_to_files(self, chunks_data: List[Dict], doc_id: str) -> List[Dict]:
        """Lưu các chunk vào file riêng biệt"""
        try:
            print(f"Đang lưu {len(chunks_data)} chunks vào files...")
            
            # Tạo thư mục document
            doc_dir = os.path.join(DATA_DIR, doc_id)
            os.makedirs(doc_dir, exist_ok=True)
            
            saved_chunks = []
            
            for i, chunk_data in enumerate(chunks_data):
                # Tạo tên file
                file_name = f"chunk_{i+1}.md"
                file_path = os.path.join(doc_dir, file_name)
                
                # Lưu nội dung chunk
                content = chunk_data.get('content', '')
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                # Cập nhật thông tin chunk (bỏ content để không lưu vào metadata)
                chunk_info = {
                    "chunk_id": chunk_data.get('chunk_id', f"{doc_id}_chunk_{i+1}"),
                    "chunk_type": chunk_data.get('chunk_type', 'article'),
                    "file_path": f"/data/{doc_id}/{file_name}",
                    "content_summary": chunk_data.get('content_summary', f'Phần {i+1} của văn bản')
                }
                
                saved_chunks.append(chunk_info)
                print(f"Đã lưu chunk {i+1}: {chunk_info['chunk_id']} ({len(content)} ký tự)")
            
            print(f"Hoàn thành lưu tất cả chunks vào {doc_dir}")
            return saved_chunks
            
        except Exception as e:
            print(f"Lỗi khi lưu chunks: {str(e)}")
            raise Exception(f"Không thể lưu chunks: {str(e)}")

    def save_metadata(self, metadata: Dict[str, Any], doc_id: str) -> str:
        """Lưu metadata vào file JSON"""
        try:
            doc_dir = os.path.join(DATA_DIR, doc_id)
            metadata_path = os.path.join(doc_dir, "metadata.json")
            
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)
            
            print(f"Đã lưu metadata tại: {metadata_path}")
            
            # Log summary metadata với auto-detection info
            print("Tóm tắt metadata đã lưu (với auto-detection):")
            print(f"  - Mã văn bản (auto): {metadata.get('doc_id')}")
            print(f"  - Loại văn bản (auto): {metadata.get('doc_type')}")
            print(f"  - Tiêu đề (auto): {metadata.get('doc_title')[:50]}...")
            print(f"  - Ngày ban hành (auto): {metadata.get('issue_date')}")
            print(f"  - Ngày hiệu lực (auto): {metadata.get('effective_date')}")
            print(f"  - Số chunks: {len(metadata.get('chunks', []))}")
            print(f"  - Văn bản liên quan: {len(metadata.get('related_documents', []))}")
            
            return metadata_path
            
        except Exception as e:
            print(f"Lỗi khi lưu metadata: {str(e)}")
            raise Exception(f"Không thể lưu metadata: {str(e)}")

    def process_document(self, file_path: str, doc_metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Xử lý toàn bộ quy trình từ tài liệu đến chunks và metadata với auto-detection"""
        try:
            doc_id = doc_metadata.get('doc_id')
            file_type = os.path.splitext(file_path)[1].lower()
            print(f"Bắt đầu xử lý tài liệu {file_type.upper()} cho document: {doc_id}")
            
            # Bước 1: Trích xuất nội dung tài liệu
            print(f"=== BƯỚC 1: TRÍCH XUẤT {file_type.upper()} ===")
            content = self.extract_document_content(file_path)
            
            # Bước 2: Gọi Gemini để phân tích, chia chunk và auto-detect metadata
            print("=== BƯỚC 2: PHÂN TÍCH VỚI GEMINI (AUTO-DETECTION) ===")
            chunked_result = self.chunk_content_with_gemini(content, doc_metadata)
            
            # Bước 3: Sử dụng auto-detected doc_id cho folder
            final_doc_id = chunked_result.get('doc_id') or doc_id
            print(f"Sử dụng doc_id: {final_doc_id} (auto-detected: {chunked_result.get('doc_id')})")
            
            # Bước 4: Lưu chunks vào files
            print("=== BƯỚC 3: LƯU CHUNKS ===")
            chunks_from_gemini = chunked_result.get('chunks', [])
            saved_chunks = self.save_chunks_to_files(chunks_from_gemini, final_doc_id)
            
            # Bước 5: Tạo metadata hoàn chỉnh với auto-detected data
            print("=== BƯỚC 4: TẠO METADATA HOÀN CHỈNH (AUTO-DETECTED) ===")
            final_metadata = {
                "doc_id": final_doc_id,
                "doc_type": chunked_result.get('doc_type', doc_metadata.get('doc_type', '')),
                "doc_title": chunked_result.get('doc_title', doc_metadata.get('doc_title', '')),
                "issue_date": chunked_result.get('issue_date', datetime.now().strftime("%d-%m-%Y")),
                "effective_date": chunked_result.get('effective_date', doc_metadata.get('effective_date', '')),
                "expiry_date": chunked_result.get('expiry_date'),
                "status": chunked_result.get('status', 'active'),
                "document_scope": chunked_result.get('document_scope', doc_metadata.get('document_scope', 'Quốc gia')),
                "replaces": chunked_result.get('replaces', []),
                "replaced_by": chunked_result.get('replaced_by'),
                "amends": chunked_result.get('amends'),
                "amended_by": chunked_result.get('amended_by'),
                "retroactive": chunked_result.get('retroactive', False),
                "retroactive_date": chunked_result.get('retroactive_date'),
                "chunks": saved_chunks,
                "related_documents": chunked_result.get('related_documents', [])
            }
            
            # Bước 6: Lưu metadata
            print("=== BƯỚC 5: LƯU METADATA ===")
            self.save_metadata(final_metadata, final_doc_id)
            
            print(f"=== HOÀN THÀNH XỬ LÝ {file_type.upper()} CHO {final_doc_id} ===")
            print(f"Kết quả (với auto-detection):")
            print(f"  - Doc ID (final): {final_doc_id}")
            print(f"  - Doc Type (auto): {final_metadata['doc_type']}")
            print(f"  - Doc Title (auto): {final_metadata['doc_title'][:50]}...")
            print(f"  - Effective Date (auto): {final_metadata['effective_date']}")
            print(f"  - Chunks: {len(saved_chunks)}")
            print(f"  - Related documents: {len(final_metadata['related_documents'])}")
            
            return {
                "doc_id": final_doc_id,
                "chunks_count": len(saved_chunks),
                "related_documents_count": len(final_metadata['related_documents']),
                "metadata": final_metadata,
                "processing_summary": f"Đã chia thành {len(saved_chunks)} chunks và tìm thấy {len(final_metadata['related_documents'])} văn bản liên quan",
                "auto_detected": {
                    "doc_id": chunked_result.get('doc_id'),
                    "doc_type": chunked_result.get('doc_type'),
                    "doc_title": chunked_result.get('doc_title'),
                    "effective_date": chunked_result.get('effective_date')
                }
            }
            
        except Exception as e:
            print(f"Lỗi trong quá trình xử lý tài liệu: {str(e)}")
            # Cleanup nếu có lỗi
            try:
                doc_dir = os.path.join(DATA_DIR, doc_metadata.get('doc_id', ''))
                if os.path.exists(doc_dir):
                    import shutil
                    shutil.rmtree(doc_dir)
                    print(f"Đã xóa thư mục lỗi: {doc_dir}")
            except:
                pass
            raise e

# Singleton instance
document_processing_service = DocumentProcessingService()