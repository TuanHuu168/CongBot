import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

class PDFExtractor:
    def __init__(self, pdf_path: str, output_docx: str):
        self.pdf_path = pdf_path
        self.output_docx = output_docx
        self.doc = Document()

    def clean_text(self, text: str) -> str:
        return ' '.join(text.split()).strip()

    def get_table_bboxes(self, tables):
        bboxes = []
        for table in tables:
            if hasattr(table, 'bbox'):
                bboxes.append(table.bbox)
        return bboxes

    def rectangles_overlap(self, rect1, rect2):
        x0_1, top_1, x1_1, bottom_1 = rect1
        x0_2, top_2, x1_2, bottom_2 = rect2
        return not (x1_1 <= x0_2 or
                    x0_1 >= x1_2 or
                    bottom_1 <= top_2 or
                    top_1 >= bottom_2)

    def group_text_by_line(self, text_words, vertical_tolerance=2):
        text_words_sorted = sorted(text_words, key=lambda w: (w['top'], w['x0']))
        lines = []
        current_line = []
        current_top = None
        for word in text_words_sorted:
            if current_top is None:
                current_top = word['top']
                current_line.append(word)
            else:
                if abs(word['top'] - current_top) <= vertical_tolerance:
                    current_line.append(word)
                else:
                    lines.append(current_line)
                    current_line = [word]
                    current_top = word['top']
        if current_line:
            lines.append(current_line)
        return lines

    def extract_content(self):
        with pdfplumber.open(self.pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                print(f"Processing page {page_num}")
                
                # Extract tables
                tables = page.extract_tables()
                table_bboxes = self.get_table_bboxes(page.find_tables())
                
                # Extract all text words
                text_words = page.extract_words()
                
                # Filter out text that falls within any table's bounding box
                filtered_text = []
                for word in text_words:
                    word_bbox = (word['x0'], word['top'], word['x1'], word['bottom'])
                    if not any(self.rectangles_overlap(word_bbox, tbbox) for tbbox in table_bboxes):
                        filtered_text.append(word)
                
                # Group text into lines
                text_lines = self.group_text_by_line(filtered_text)
                
                # Prepare text elements
                text_elements = []
                for line in text_lines:
                    text_content = ' '.join([w['text'] for w in line])
                    text_content = self.clean_text(text_content)
                    if text_content:
                        text_elements.append({'content': text_content, 'top': line[0]['top']})
                
                # Prepare table elements
                table_elements = []
                for table in tables:
                    if table:
                        # Check for varying column counts and split tables if necessary
                        current_table = []
                        previous_cols = len(table[0]) if table else 0
                        current_row = []
                        for row in table:
                            if len(row) != previous_cols:
                                if current_row:
                                    current_table.append(current_row)
                                current_row = [row]
                                previous_cols = len(row)
                            else:
                                current_row.append(row)
                        if current_row:
                            current_table.append(current_row)
                        # Add each sub-table as a separate table element
                        for sub_table in current_table:
                            table_elements.append({'content': sub_table, 'top': 0})  # Top position not available in extracted tables
                
                # Combine and sort all elements
                all_elements = text_elements + table_elements
                sorted_elements = sorted(all_elements, key=lambda x: x['top'])
                
                # Write elements into the Word document
                for element in sorted_elements:
                    if isinstance(element['content'], list):
                        self.write_table(element['content'])
                    else:
                        self.write_text(element['content'])
                
                # Add a space after each page
                self.doc.add_paragraph()
        
        self.doc.save(self.output_docx)
        print(f"Saved to {self.output_docx}")

    def write_table(self, table_data):
        if not table_data:
            return
        
        # Determine the number of columns
        max_cols = max(len(row) for row in table_data)
        table = self.doc.add_table(rows=len(table_data), cols=max_cols)
        table.style = 'Table Grid'
        
        for row_idx, row in enumerate(table_data):
            for col_idx, cell in enumerate(row):
                # Clean cell text
                cell_text = self.clean_text(str(cell).replace('\n', ' '))
                table.cell(row_idx, col_idx).text = cell_text
                # Set alignment
                if row_idx == 0:
                    table.cell(row_idx, col_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    table.cell(row_idx, col_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        self.doc.add_paragraph()

    def write_text(self, text):
        if text := self.clean_text(text):
            para = self.doc.add_paragraph(text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Usage
pdf_path = r"D:\University\DATN\PDF\PL QD 1072 ngay 27-2-2024.pdf"
output_docx = "output.docx"

extractor = PDFExtractor(pdf_path, output_docx)
extractor.extract_content()